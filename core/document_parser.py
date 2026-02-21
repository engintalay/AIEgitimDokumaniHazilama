"""Document parser for PDF, DOCX, and TXT files."""
import os
from typing import Optional
import fitz  # pymupdf
from docx import Document


class DocumentParser:
    """Parse various document formats and extract text."""
    
    @staticmethod
    def parse(file_path: str, mode: str = 'paragraph') -> str:
        """Parse document and return text content."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            return DocumentParser._parse_pdf(file_path, mode)
        elif ext in ['.docx', '.doc']:
            return DocumentParser._parse_docx(file_path)
        elif ext == '.txt':
            return DocumentParser._parse_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    @staticmethod
    def _parse_pdf(file_path: str, mode: str = 'paragraph') -> str:
        """Extract text, tables, and images from PDF using pymupdf."""
        text_blocks = []
        try:
            doc = fitz.open(file_path)
            basename = os.path.splitext(os.path.basename(file_path))[0]
            img_dir = os.path.join('data', 'images', basename)
            os.makedirs(img_dir, exist_ok=True)
            
            for page_index in range(len(doc)):
                page = doc[page_index]
                page_content = []
                
                # Add page marker if in page mode
                if mode == 'page':
                    page_content.append(f"--- SAYFA {page_index + 1} ---")
                
                # 1. Extract tables first
                tabs = page.find_tables()
                table_areas = [t.bbox for t in tabs.tables]
                
                # 2. Extract text blocks - using default extraction flags
                # Default flags preserve the best mapping for custom Turkish fonts
                blocks = page.get_text("blocks", sort=True)
                for b in blocks:
                    block_bbox = b[:4]
                    is_inside_table = False
                    for t_bbox in table_areas:
                        # Check if block center is inside table bbox
                        mid_x = (block_bbox[0] + block_bbox[2]) / 2
                        mid_y = (block_bbox[1] + block_bbox[3]) / 2
                        if (t_bbox[0] <= mid_x <= t_bbox[2] and 
                            t_bbox[1] <= mid_y <= t_bbox[3]):
                            is_inside_table = True
                            break
                    
                    if not is_inside_table:
                        content = b[4].strip()
                        if content:
                            # Normalize whitespace but keep Turkish characters intact
                            page_content.append(content)
                
                # 3. Add extracted tables as Markdown
                for tab in tabs.tables:
                    df = tab.to_pandas()
                    if not df.empty:
                        md_table = "\n\n" + df.to_markdown(index=False) + "\n\n"
                        page_content.append(md_table)
                
                # 4. Extract images
                image_list = page.get_images(full=True)
                if image_list:
                    for img_index, img in enumerate(image_list, 1):
                        xref = img[0]
                        base_image = doc.extract_image(xref)
                        image_bytes = base_image["image"]
                        image_ext = base_image["ext"]
                        img_filename = f"image_p{page_index+1}_n{img_index}.{image_ext}"
                        img_path = os.path.join(img_dir, img_filename)
                        
                        with open(img_path, "wb") as f:
                            f.write(image_bytes)
                        
                        marker = f"\n\n[GÖRSEL: data/images/{basename}/{img_filename}]\n\n"
                        page_content.append(marker)
                
                # Join page content and add to main list
                if page_content:
                    text_blocks.append("\n\n".join(page_content))
                
            doc.close()
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")
        
        # If page mode, we might want a different joiner, but \n\n is safe.
        return "\n\n".join(text_blocks)
    
    @staticmethod
    def _parse_docx(file_path: str) -> str:
        """Extract text from DOCX."""
        doc = Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    
    @staticmethod
    def _parse_txt(file_path: str) -> str:
        """Read text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
